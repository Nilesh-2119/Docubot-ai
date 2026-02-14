"""
File parser utility.
Extracts text from PDF, DOCX, XLSX, CSV, and TXT files.
Spreadsheet files (XLSX, CSV) are parsed with column headers
attached to every row for better RAG retrieval.
"""
import csv
import io


def extract_text(file_path: str, file_type: str) -> str:
    """Extract text content from a file based on its type."""
    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "xlsx": _extract_xlsx,
        "csv": _extract_csv,
        "txt": _extract_txt,
    }

    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor(file_path)


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    text_parts = []

    for page_num, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

    return "\n\n".join(text_parts)


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _format_row_with_headers(headers: list[str], row_values: list[str], row_num: int) -> str:
    """Format a spreadsheet row as 'Column: Value' pairs for better RAG retrieval."""
    pairs = []
    for header, value in zip(headers, row_values):
        value = value.strip() if value else ""
        if value:
            pairs.append(f"{header}: {value}")
    if not pairs:
        return ""
    return f"[Row {row_num}] " + " | ".join(pairs)


def _extract_xlsx(file_path: str) -> str:
    """Extract text from XLSX spreadsheet with column headers on every row."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"[Sheet: {sheet_name}]")

        headers = []
        for row_num, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            row_values = [str(cell) if cell is not None else "" for cell in row]

            # First row = headers
            if row_num == 1:
                headers = [v.strip() if v.strip() else f"Column{i+1}" for i, v in enumerate(row_values)]
                text_parts.append("Columns: " + " | ".join(headers))
                continue

            # Data rows — attach headers to each value
            if headers:
                formatted = _format_row_with_headers(headers, row_values, row_num)
            else:
                formatted = " | ".join(v for v in row_values if v)

            if formatted:
                text_parts.append(formatted)

    wb.close()
    return "\n".join(text_parts)


def _extract_csv(file_path: str) -> str:
    """Extract text from CSV file with column headers on every row."""
    text_parts = []

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        headers = []

        for row_num, row in enumerate(reader, start=1):
            row_values = [cell.strip() for cell in row]

            # First row = headers
            if row_num == 1:
                headers = [v if v else f"Column{i+1}" for i, v in enumerate(row_values)]
                text_parts.append("Columns: " + " | ".join(headers))
                continue

            # Data rows — attach headers
            if headers:
                formatted = _format_row_with_headers(headers, row_values, row_num)
            else:
                formatted = " | ".join(v for v in row_values if v)

            if formatted:
                text_parts.append(formatted)

    return "\n".join(text_parts)


def _extract_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
